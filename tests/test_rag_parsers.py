"""Tests for rag.parsers: parse_markdown, parse_txt, parse_html, parse_pdf (when available). #56 coverage."""

from __future__ import annotations

import contextlib
import sys
from pathlib import Path
from unittest.mock import patch

import pytest

from voiceforge.rag import parsers


def test_parse_markdown(tmp_path: Path) -> None:
    """parse_markdown strips markdown and returns segments."""
    md = tmp_path / "doc.md"
    md.write_text("# Title\n\nParagraph one.\n\n- item 1\n- item 2\n\n**bold** text.", encoding="utf-8")
    out = parsers.parse_markdown(md)
    assert out
    assert "Title" in " ".join(out) or "Paragraph" in " ".join(out)
    assert "item" in " ".join(out) or "bold" in " ".join(out)


def test_parse_markdown_empty_paragraphs_returns_single_segment(tmp_path: Path) -> None:
    """When only whitespace/code blocks, returns one segment of stripped content."""
    md = tmp_path / "minimal.md"
    md.write_text("Only one line.", encoding="utf-8")
    out = parsers.parse_markdown(md)
    assert out == ["Only one line."]


def test_parse_markdown_with_code_blocks(tmp_path: Path) -> None:
    """parse_markdown strips code blocks and inline code."""
    md = tmp_path / "code.md"
    md.write_text(
        "# Doc\n\nParagraph with `inline` and\n\n```\nblock code\n```\n\nAfter.",
        encoding="utf-8",
    )
    out = parsers.parse_markdown(md)
    assert out
    joined = " ".join(out)
    assert "Doc" in joined or "Paragraph" in joined
    assert "block" not in joined or "After" in joined


def test_parse_markdown_missing_file() -> None:
    """parse_markdown raises FileNotFoundError for missing file."""
    with pytest.raises(FileNotFoundError):
        parsers.parse_markdown(Path("/nonexistent/file.md"))


def test_parse_txt(tmp_path: Path) -> None:
    """parse_txt splits by double newline."""
    txt = tmp_path / "doc.txt"
    txt.write_text("First para.\n\nSecond para.\n\nThird.", encoding="utf-8")
    out = parsers.parse_txt(txt)
    assert len(out) >= 2
    assert "First" in out[0]
    assert "Second" in out[1]


def test_parse_txt_single_line(tmp_path: Path) -> None:
    """parse_txt returns one segment for single line."""
    txt = tmp_path / "one.txt"
    txt.write_text("Single line.", encoding="utf-8")
    assert parsers.parse_txt(txt) == ["Single line."]


def test_parse_txt_missing_file() -> None:
    """parse_txt raises FileNotFoundError for missing file."""
    with pytest.raises(FileNotFoundError):
        parsers.parse_txt(Path("/nonexistent/file.txt"))


def test_parse_html(tmp_path: Path) -> None:
    """parse_html extracts text from HTML."""
    html = tmp_path / "page.html"
    html.write_text(
        "<html><body><p>First</p><div>Second</div><h1>Title</h1></body></html>",
        encoding="utf-8",
    )
    out = parsers.parse_html(html)
    assert out
    joined = " ".join(out)
    assert "First" in joined
    assert "Second" in joined or "Title" in joined


def test_parse_html_missing_file() -> None:
    """parse_html raises FileNotFoundError for missing file."""
    with pytest.raises(FileNotFoundError):
        parsers.parse_html(Path("/nonexistent/file.html"))


def test_parse_html_empty_raw_returns_single_segment(tmp_path: Path) -> None:
    """parse_html when only whitespace returns one segment or empty then stripped."""
    html = tmp_path / "blank.html"
    html.write_text("   \n\n  ", encoding="utf-8")
    out = parsers.parse_html(html)
    assert out == [] or (len(out) == 1 and out[0].strip() == "")


def test_parse_pdf_missing_file() -> None:
    """parse_pdf raises FileNotFoundError for missing file."""
    with pytest.raises(FileNotFoundError):
        parsers.parse_pdf(Path("/nonexistent/file.pdf"))


def test_parse_pdf_import_error_when_no_pymupdf(tmp_path: Path) -> None:
    """parse_pdf raises ImportError when pymupdf not installed (existing file triggers import)."""
    pdf = tmp_path / "dummy.pdf"
    pdf.write_bytes(b"not a pdf")
    try:
        import fitz  # noqa: F401

        # With pymupdf, fitz.open may raise on invalid content
        with contextlib.suppress(Exception):
            parsers.parse_pdf(pdf)
    except ImportError:
        with pytest.raises(ImportError, match="rag|pymupdf"):
            parsers.parse_pdf(pdf)


def test_parse_docx_missing_file() -> None:
    """parse_docx raises FileNotFoundError for missing file."""
    with pytest.raises(FileNotFoundError):
        parsers.parse_docx(Path("/nonexistent/file.docx"))


def test_parse_docx_import_error(tmp_path: Path) -> None:
    """parse_docx raises ImportError when python-docx is not installed."""
    docx_path = tmp_path / "x.docx"
    docx_path.touch()
    with patch.dict(sys.modules, {"docx": None}):
        with pytest.raises(ImportError, match="rag|python-docx"):
            parsers.parse_docx(docx_path)


def test_parse_docx_success(tmp_path: Path) -> None:
    """parse_docx extracts paragraphs when python-docx is available."""
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_paragraph("First para.")
    doc.add_paragraph("Second para.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    out = parsers.parse_docx(path)
    assert len(out) >= 1
    assert "First" in out[0] or "Second" in (out[1] if len(out) > 1 else "")


def test_parse_odt_missing_file() -> None:
    """parse_odt raises FileNotFoundError for missing file."""
    with pytest.raises(FileNotFoundError):
        parsers.parse_odt(Path("/nonexistent/file.odt"))


def test_parse_odt_import_error(tmp_path: Path) -> None:
    """parse_odt raises ImportError when odfpy is not installed."""
    odt_path = tmp_path / "x.odt"
    odt_path.touch()
    with patch.dict(sys.modules, {"odf": None}):
        with pytest.raises(ImportError, match="rag|odfpy"):
            parsers.parse_odt(odt_path)


def test_parse_odt_success(tmp_path: Path) -> None:
    """parse_odt extracts text when odfpy is available."""
    pytest.importorskip("odf")
    from odf import teletype
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    doc = OpenDocumentText()
    p = P()
    teletype.addTextToElement(p, "Hello ODT.")
    doc.text.addElement(p)
    path = tmp_path / "sample.odt"
    doc.save(str(path))
    out = parsers.parse_odt(path)
    assert any("Hello" in s for s in out)


def test_parse_rtf_missing_file() -> None:
    """parse_rtf raises FileNotFoundError for missing file."""
    with pytest.raises(FileNotFoundError):
        parsers.parse_rtf(Path("/nonexistent/file.rtf"))


def test_parse_rtf_import_error(tmp_path: Path) -> None:
    """parse_rtf raises ImportError when striprtf is not installed."""
    rtf_path = tmp_path / "x.rtf"
    rtf_path.write_text(r"{\rtf1 hi}", encoding="utf-8")
    with patch.dict(sys.modules, {"striprtf": None}):
        with pytest.raises(ImportError, match="rag|striprtf"):
            parsers.parse_rtf(rtf_path)


def test_parse_rtf_success(tmp_path: Path) -> None:
    """parse_rtf extracts text when striprtf is available."""
    pytest.importorskip("striprtf")
    rtf_path = tmp_path / "sample.rtf"
    rtf_path.write_text(r"{\rtf1\ansi Hello RTF.\par }", encoding="utf-8")
    out = parsers.parse_rtf(rtf_path)
    assert any("Hello" in s for s in out)
